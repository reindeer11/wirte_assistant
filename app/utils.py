from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
import re

def format_content(content: str) -> str:
    """
    Post-process content to ensure proper formatting:
    - First non-empty line is treated as title
    - Remove any markdown symbols if present
    """
    lines = content.strip().split('\n')
    formatted_lines = []
    
    for line in lines:
        # Remove markdown heading symbols
        cleaned = re.sub(r'^#{1,6}\s*', '', line)
        # Remove markdown bold/italic
        cleaned = re.sub(r'\*{1,2}([^*]+)\*{1,2}', r'\1', cleaned)
        formatted_lines.append(cleaned)
    
    return '\n'.join(formatted_lines)

def create_docx(content: str, filename: str = "output.docx") -> str:
    """
    Creates a docx file from content.
    - First non-empty line is centered as title
    - Subsequent content is formatted with proper paragraph indentation
    Returns the absolute path to the generated file.
    """
    doc = Document()
    
    # Post-process content to remove markdown
    content = format_content(content)
    
    lines = content.split('\n')
    is_first_content = True
    
    for i, line in enumerate(lines):
        stripped = line.strip()
        
        if not stripped:
            # Empty line - skip
            continue
        
        if is_first_content:
            # First non-empty line is the title - center it
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(stripped)
            title_run.bold = True
            title_run.font.size = Pt(18)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.space_after = Pt(24)
            is_first_content = False
        elif re.match(r'^\d+(\.\d+)*[\.、\s]', stripped):
            # Numbered subheading (e.g., 1.1, 2.1)
            heading_para = doc.add_paragraph()
            heading_run = heading_para.add_run(stripped)
            heading_run.bold = True
            heading_run.font.size = Pt(14)
            heading_para.space_before = Pt(16)
            heading_para.space_after = Pt(8)
        else:
            # Regular paragraph with 2-character indent
            para = doc.add_paragraph()
            # Add indentation (2 Chinese characters ≈ 0.33 inch)
            para.paragraph_format.first_line_indent = Inches(0.33)
            para.add_run(stripped)
            para.paragraph_format.line_spacing = 1.5

    # Ensure output directory exists (static/downloads)
    output_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static", "downloads")
    os.makedirs(output_dir, exist_ok=True)
    
    file_path = os.path.join(output_dir, filename)
    doc.save(file_path)
    return file_path
